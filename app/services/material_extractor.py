from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_EXTRACTED_CHARS = 12_000
MAX_PDF_PAGES = 30
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class MaterialExtractionError(ValueError):
    pass


def _normalize_text(value: str) -> str:
    lines = [" ".join(line.split()) for line in value.splitlines() if line.strip()]
    return "\n".join(lines).strip()


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise MaterialExtractionError("文本编码无法识别，请另存为 UTF-8 后重试。")


def _extract_docx(data: bytes) -> str:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise MaterialExtractionError("Word 文档无法读取，请确认文件未损坏且为 .docx 格式。") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def _extract_pdf(data: bytes) -> tuple[str, bool]:
    try:
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            raise MaterialExtractionError("暂不支持加密 PDF，请解密后重试。")
        pages = reader.pages[:MAX_PDF_PAGES]
        text = "\n".join(page.extract_text() or "" for page in pages)
        return text, len(reader.pages) > MAX_PDF_PAGES
    except MaterialExtractionError:
        raise
    except Exception as exc:
        raise MaterialExtractionError("PDF 无法读取，请确认文件未损坏。") from exc


def extract_material_text(file_name: str, data: bytes) -> tuple[str, bool]:
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise MaterialExtractionError("仅支持 PDF、Word (.docx)、Markdown 和 TXT 文档。")

    page_truncated = False
    if suffix in {".txt", ".md"}:
        raw_text = _decode_text(data)
    elif suffix == ".docx":
        raw_text = _extract_docx(data)
    else:
        raw_text, page_truncated = _extract_pdf(data)

    text = _normalize_text(raw_text)
    if not text:
        if suffix == ".pdf":
            raise MaterialExtractionError("PDF 中没有可复制文本；当前版本暂不支持扫描件 OCR。")
        raise MaterialExtractionError("文档中没有可提取的文字。")

    truncated = page_truncated or len(text) > MAX_EXTRACTED_CHARS
    return text[:MAX_EXTRACTED_CHARS], truncated
