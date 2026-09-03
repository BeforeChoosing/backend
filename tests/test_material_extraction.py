from io import BytesIO

from docx import Document

from app.api import profile as profile_api
from app.schemas.multimodal import MultimodalEvidenceItem, MultimodalEvidenceResponse


def test_extract_text_material(authenticated_client) -> None:
    response = authenticated_client.post(
        "/api/v1/profile/materials/extract",
        files={"file": ("experience.md", "我负责用户访谈，并根据反馈修改原型。".encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 200
    assert response.json()["text"] == "我负责用户访谈，并根据反馈修改原型。"
    assert response.json()["truncated"] is False


def test_extract_docx_paragraphs_and_tables(authenticated_client) -> None:
    document = Document()
    document.add_paragraph("项目经历：完成需求分析。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "行动"
    table.cell(0, 1).text = "访谈用户"
    buffer = BytesIO()
    document.save(buffer)

    response = authenticated_client.post(
        "/api/v1/profile/materials/extract",
        files={"file": ("resume.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    assert "项目经历" in response.json()["text"]
    assert "行动 | 访谈用户" in response.json()["text"]


def test_reject_unsupported_material(authenticated_client) -> None:
    response = authenticated_client.post(
        "/api/v1/profile/materials/extract",
        files={"file": ("legacy.doc", b"not-a-word-file", "application/msword")},
    )

    assert response.status_code == 422
    assert "仅支持" in response.json()["detail"]


def test_multimodal_extract_returns_candidate_region(monkeypatch, authenticated_client) -> None:
    class FakeExtractor:
        async def extract(self, *, file_name, data, mime_type):
            return MultimodalEvidenceResponse(
                file_name=file_name,
                file_sha256="a" * 64,
                mime_type=mime_type or "image/png",
                page_count=1,
                model="qwen-vl-ocr",
                items=[
                    MultimodalEvidenceItem(
                        id="mm:test:1:1",
                        source_ref="material:test:page:1:region:1",
                        page=1,
                        bbox=[10, 20, 300, 400],
                        label="项目行动",
                        quote="我负责了用户访谈",
                        confidence=0.9,
                    )
                ],
            )

    monkeypatch.setattr(profile_api, "MultimodalEvidenceExtractor", FakeExtractor)
    response = authenticated_client.post(
        "/api/v1/profile/materials/multimodal-extract",
        files={"file": ("evidence.png", b"image", "image/png")},
    )

    assert response.status_code == 200
    assert response.json()["items"][0]["status"] == "candidate"
    assert response.json()["items"][0]["page"] == 1
